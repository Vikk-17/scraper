import sys
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document
import json
import os
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()


class EmailSender:
    """
    To send the email with a document of vulnerabilities
    :params: raw_data, contains the user email and vulnerability details
    :output: generate a report and send it via email
    """

    def __init__(self, raw_data: str):
        self.sender_email = "cyberalerter11@gmail.com"
        self.sender_password = "sbdv cfei kzuw ajtv"  # Use environment variable for password
        self.subject = "Mail regarding critical vulnerabilities"
        self.body = "Here is your report of vulnerabilities:\n"
        try:
            self.obj_data = json.loads(raw_data)
        except Exception as e:
            raise Exception(f"Error while processing the raw data: {e}")

    def get_email(self) -> str:
        """Get the email from the dictionary file"""
        try:
            return self.obj_data.get("userEmail")
        except Exception as e:
            raise Exception(f"Can't get the user email: {e}")

    def format_data(self) -> str:
        """Get the description and format it"""
        try:
            scan_details_obj = self.obj_data.get("scanDetails", {})
            results = scan_details_obj.get("results", [])

            output_data = f"Product Name: {scan_details_obj.get('productName', 'Unknown')}\n"
            output_data += f"Product Version: {scan_details_obj.get('productVersion', 'N/A')}\n"
            output_data += "\nVulnerabilities:\n"
            output_data += "-" * 70 + "\n"

            for vuln in results:
                output_data += f"""
                    CVE ID: {vuln.get('cve_id')}
                    Severity: {vuln.get('baseSeverity')}
                    Description: {vuln.get('vulnerabilityDescription')}
                    Mitigation: {vuln.get('Mitigation', 'N/A')}
                    Published Date: {vuln.get('published_date')}
                    Last Modified: {vuln.get('last_modified')}
                    URL: {vuln.get('oemUrl')}
                        """
                output_data += "-" * 70 + "\n"

            return output_data

        except Exception as e:
            raise Exception(f"Can't create the format: {e}")

    def create_doc(self) -> str:
        """To create the document"""
        try:
            doc = Document()
            doc.add_heading("Details of your product", level=1)
            doc.add_paragraph("Below you may find the vulnerabilities for your product:\n")
            formatted_data = self.format_data()  # Now contains data for all vulnerabilities
            doc.add_paragraph(formatted_data)
            filename = "vuln_report.docx"
            doc.save(filename)
            return filename
        except Exception as e:
            raise Exception(f"Can't generate or save file: {e}")

    def send_email(self):
        """To send the email with the attachment"""
        try:
            message = MIMEMultipart()
            message["From"] = self.sender_email
            message["To"] = self.get_email()
            message["Subject"] = self.subject

            # Email body
            message.attach(MIMEText(self.body, "plain"))
            file_to_send = self.create_doc()

            # Attach the file
            with open(file_to_send, "rb") as attachment_file:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment_file.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={file_to_send}")
            message.attach(part)

            # Send the email
            with smtplib.SMTP("smtp.gmail.com", 587) as server:
                server.starttls()
                server.login(self.sender_email, self.sender_password)
                server.send_message(message)
                print("Email sent successfully!")

        except Exception as e:
            print(f"Error in sending email: {e}")



def main():
    try:
        
        # with open("input_data.json", 'r') as file:
        #    data = json.load(file)
        input_raw_data: str = sys.stdin.read()
    except Exception as e:
        return f"Can't access data from backend {e}"

    # except json.JSONDecodeError as e:
    #     print(f"Error in JSON format: {e}")
    #     return  # Stop further execution if JSON is invalid

    try:
        send_email_instance = EmailSender(input_raw_data)
        send_email_instance.send_email()
    except Exception as e:
        print(f"Unexpected error while sending email: {e}")


if __name__ == "__main__":
    main()
